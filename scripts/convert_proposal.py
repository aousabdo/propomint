"""Convert a proposal text file into cleaned Markdown and Word documents.

This utility removes the box-drawing artifacts ("|", box borders, long
hyphen dividers, etc.) that appear in the generated `.txt` export before
writing:

* `<basename>_clean.md`
* `<basename>_clean.docx`

Usage
-----
python scripts/convert_proposal.py proposal_pretty_full.txt

An optional ``--output-dir`` argument lets you direct the results to a
different folder.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable, List

try:
    from docx import Document  # type: ignore
except ImportError as exc:  # pragma: no cover - surfaced at runtime
    raise SystemExit(
        "The python-docx package is required. Install it with `pip install python-docx`."
    ) from exc


UNWANTED_CHARS = "┏┓┗┛┳┻┯┷┴┬┠┨┰┸┞┟┢┡┥┤┬┼┮┾┃│─━═╍╎╏╱╲╳╴╵╶╷╾╿|"
TRANSLATION_TABLE = str.maketrans("", "", UNWANTED_CHARS)
SEPARATOR_PATTERN = re.compile(r"^[-=─━\s]+$")


def clean_lines(lines: Iterable[str]) -> List[str]:
    cleaned: List[str] = []
    for raw in lines:
        text = raw.rstrip("\n").translate(TRANSLATION_TABLE).strip()
        if not text:
            cleaned.append("")
            continue
        if SEPARATOR_PATTERN.match(text):
            cleaned.append("")
            continue
        cleaned.append(text)
    # Collapse multiple blank lines into single blanks for readability.
    normalized: List[str] = []
    previous_blank = False
    for line in cleaned:
        is_blank = line == ""
        if is_blank and previous_blank:
            continue
        normalized.append(line)
        previous_blank = is_blank
    return normalized


def write_markdown(lines: Iterable[str], destination: Path) -> None:
    destination.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_word_doc(lines: Iterable[str], destination: Path) -> None:
    document = Document()
    for line in lines:
        if line:
            document.add_paragraph(line)
        else:
            document.add_paragraph("")
    document.save(destination)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Path to the source .txt proposal file")
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="Optional directory for the generated files (defaults to the input file's parent)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path: Path = args.input.expanduser().resolve()
    if not input_path.exists():
        raise SystemExit(f"Input file not found: {input_path}")

    output_dir = (args.output_dir or input_path.parent).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    lines = input_path.read_text(encoding="utf-8", errors="ignore").splitlines()
    cleaned_lines = clean_lines(lines)

    base_name = input_path.stem
    md_path = output_dir / f"{base_name}_clean.md"
    docx_path = output_dir / f"{base_name}_clean.docx"

    write_markdown(cleaned_lines, md_path)
    write_word_doc(cleaned_lines, docx_path)

    print(f"Markdown written to: {md_path}")
    print(f"Word document written to: {docx_path}")


if __name__ == "__main__":
    main()

